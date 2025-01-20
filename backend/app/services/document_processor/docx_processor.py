# app/services/document_processor/docx_processor.py

from .base import BaseDocumentProcessor
from typing import Dict, Any
import logging
from docx import Document
from datetime import datetime

logger = logging.getLogger(__name__)


class DOCXProcessor(BaseDocumentProcessor):
    """
    Processor for DOCX documents using python-docx.
    """

    def extract_text(self, file_path: str) -> str:
        """
        Extract text content from a DOCX file.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Extracted text content

        Raises:
            Exception: If DOCX processing fails
        """
        try:
            logger.info(f"Starting text extraction from DOCX: {file_path}")

            doc = Document(file_path)
            text_content = []

            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():  # Only add non-empty paragraphs
                    text_content.append(para.text)

            # Join all text with double newlines between paragraphs
            full_text = "\n\n".join(text_content)

            logger.info(f"Successfully extracted {len(doc.paragraphs)} paragraphs")
            return full_text

        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise

    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from a DOCX file.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Dictionary containing DOCX metadata

        Raises:
            Exception: If metadata extraction fails
        """
        try:
            logger.info(f"Extracting metadata from DOCX: {file_path}")

            doc = Document(file_path)
            core_properties = doc.core_properties

            metadata = {
                "title": core_properties.title or "",
                "author": core_properties.author or "",
                "subject": core_properties.subject or "",
                "created": (
                    core_properties.created.isoformat()
                    if core_properties.created
                    else ""
                ),
                "modified": (
                    core_properties.modified.isoformat()
                    if core_properties.modified
                    else ""
                ),
                "last_modified_by": core_properties.last_modified_by or "",
                "revision": core_properties.revision or 1,
                "file_type": "docx",
                "paragraph_count": len(doc.paragraphs),
            }

            logger.info("Successfully extracted DOCX metadata")
            logger.debug(f"Extracted metadata: {metadata}")

            return metadata

        except Exception as e:
            logger.error(f"Error extracting DOCX metadata: {str(e)}")
            raise
